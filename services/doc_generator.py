from docx import Document
import uuid
import os

def generate_doc(user_request, task_outputs):

    os.makedirs(
        "generated_docs",
        exist_ok=True
    )

    filename = f"generated_docs/{uuid.uuid4()}.docx"

    doc = Document()

    doc.add_heading(
        "Business Proposal",
        level=1
    )

    doc.add_heading(
        "User Request",
        level=2
    )

    doc.add_paragraph(
        user_request
    )

    doc.add_heading(
        "Generated Content",
        level=2
    )

    for item in task_outputs:

        doc.add_heading(
            item["task"],
            level=3
        )

        doc.add_paragraph(
            item["result"]
        )

    doc.save(
        filename
    )

    return filename